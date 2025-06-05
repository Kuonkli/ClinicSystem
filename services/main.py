import os
import uuid
from flask import Flask, request, send_file, jsonify, Response
import img2pdf
from PyPDF2 import PdfMerger
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import json
from datetime import datetime

app = Flask(__name__)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
STORAGE_DIR = os.path.abspath(os.path.join(BASE_DIR, '..', 'storage', 'documents'))
UPLOAD_DIR = os.path.join(STORAGE_DIR, 'temp_uploads')
OUTPUT_DIR = os.path.join(STORAGE_DIR, 'merged')

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in {'pdf', 'jpg', 'jpeg', 'png'}


@app.route('/api/merge', methods=['POST'])
def merge_documents():
    if 'files' not in request.files:
        return jsonify({"error": "No files provided"}), 400

    files = request.files.getlist('files')
    if not files or all(file.filename == '' for file in files):
        return jsonify({"error": "No files selected"}), 400

    saved_files = []
    pdf_files = []

    try:
        for file in files:
            if not file or file.filename == '':
                continue

            if not allowed_file(file.filename):
                return jsonify({"error": f"Invalid file type: {file.filename}"}), 400

            filename = secure_filename(file.filename)
            file_path = os.path.join(UPLOAD_DIR, filename)
            file.save(file_path)
            saved_files.append(file_path)

            # Convert images to PDF if needed
            if file_path.lower().endswith(('.jpg', '.jpeg', '.png')):
                pdf_path = f"{file_path}.pdf"
                try:
                    with open(pdf_path, "wb") as f:
                        f.write(img2pdf.convert(file_path))
                    pdf_files.append(pdf_path)
                except Exception as e:
                    return jsonify({"error": f"Failed to convert image to PDF: {str(e)}"}), 500
            else:
                pdf_files.append(file_path)

        if not pdf_files:
            return jsonify({"error": "No valid files to merge"}), 400

        merged_filename = f"merged_{uuid.uuid4().hex}.pdf"
        merged_path = os.path.join(OUTPUT_DIR, merged_filename)

        merger = PdfMerger()
        for pdf in pdf_files:
            try:
                merger.append(pdf)
            except Exception as e:
                return jsonify({"error": f"Failed to merge PDF: {str(e)}"}), 500

        merger.write(merged_path)
        merger.close()

        # Clean up - only delete files that exist
        for file_path in saved_files:
            if os.path.exists(file_path):
                try:
                    os.remove(file_path)
                except:
                    pass

        for pdf_path in pdf_files:
            if pdf_path != merged_path and os.path.exists(pdf_path):
                try:
                    os.remove(pdf_path)
                except:
                    pass

        return send_file(
            merged_path,
            as_attachment=True,
            download_name=merged_filename,
            mimetype='application/pdf'
        )

    except Exception as e:
        # Clean up in case of any error
        for file_path in saved_files:
            if os.path.exists(file_path):
                try:
                    os.remove(file_path)
                except:
                    pass
        for pdf_path in pdf_files:
            if os.path.exists(pdf_path):
                try:
                    os.remove(pdf_path)
                except:
                    pass
        return jsonify({"error": f"Internal server error: {str(e)}"}), 500


def validate_and_parse_data(data):
    try:
        # Проверка обязательных полей
        required_fields = {
            'appointment': ['ID', 'appointment_date'],
            'doctor': ['User'],
            'doctor/User': ['last_name', 'name'],
            'service': ['name']
        }

        for section, fields in required_fields.items():
            if '/' in section:
                parts = section.split('/')
                current = data
                for part in parts:
                    if part not in current:
                        raise ValueError(f"Missing section: {section}")
                    current = current[part]
            elif section not in data:
                raise ValueError(f"Missing section: {section}")

        date_str = data['appointment']['appointment_date']
        try:
            appointment_date = datetime.strptime(date_str, "%Y-%m-%dT%H:%M:%S%z")
        except ValueError:
            appointment_date = datetime.strptime(date_str, "%Y-%m-%dT%H:%M:%S")

        return {
            'appointment_id': data['appointment']['ID'],
            'appointment_date': appointment_date,
            'start_time': data['appointment'].get('start_time', ''),
            'doctor_name': f"{data['doctor']['User']['last_name']} {data['doctor']['User']['name']}",
            'doctor_specialty': data['doctor'].get('specialty', ''),
            'cabinet': data['doctor'].get('cabinet', ''),
            'service_name': data['service']['name'],
            'price': data['service'].get('price', 0)
        }

    except Exception as e:
        raise ValueError(f"Invalid data format: {str(e)}")


def generate_ticket(parsed_data):
    doc = Document()

    # Настройка стилей
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # Заголовок
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Талон на прием')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 0, 0)

    # Основная информация
    doc.add_paragraph("Уважаемый пациент, ваш талон на прием:", style='Intense Quote')

    # Создаем таблицу с данными
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'

    for row in table.rows:
        row.height = Pt(28)
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1.5  # 1.5 интервала
                paragraph.paragraph_format.space_before = Pt(4)  # Отступ сверху
                paragraph.paragraph_format.space_after = Pt(4)  # Отступ снизу

    rows_data = [
        ("Дата приема:", parsed_data['appointment_date'].strftime("%d.%m.%Y")),
        ("Время приема:", parsed_data['start_time']),
        ("Врач:", parsed_data['doctor_name']),
        ("Специальность:", parsed_data['doctor_specialty']),
        ("Кабинет:", parsed_data['cabinet']),
        ("Услуга:", parsed_data['service_name']),
        ("Стоимость:", f"{parsed_data['price']} руб." if parsed_data['price'] else "Бесплатно")
    ]

    for i, (label, value) in enumerate(rows_data):
        left_cell = table.cell(i, 0)
        left_cell.text = label
        left_cell.paragraphs[0].runs[0].font.bold = True
        left_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        right_cell = table.cell(i, 1)
        right_cell.text = value
        right_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph("\nПожалуйста, приходите за 10 минут до назначенного времени.", style='Intense Quote')

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


@app.route('/generate/ticket', methods=['POST'])
def handle_generate_ticket():
    try:
        if not request.is_json:
            return jsonify({"error": "Request must be JSON"}), 400

        data = request.get_json()
        if not data:
            return jsonify({"error": "No data provided"}), 400

        parsed_data = validate_and_parse_data(data)

        ticket_stream = generate_ticket(parsed_data)
        file_content = ticket_stream.getvalue()
        ticket_stream.close()

        return Response(
            file_content,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={
                'Content-Disposition': f'attachment; filename=ticket_{parsed_data["appointment_id"]}.docx'
            }
        )

    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        return jsonify({"error": f"Internal server error: {str(e)}"}), 500


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)